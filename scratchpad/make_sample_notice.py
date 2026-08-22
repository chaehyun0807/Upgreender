from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("2026학년도 AI 창업 아이디어 공모전 안내", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "학생 여러분의 창의적인 아이디어를 응원합니다. 2026학년도 AI 창업 아이디어 공모전을 아래와 같이 "
    "개최하오니 관심 있는 학생들의 많은 참여 바랍니다."
)

doc.add_heading("1. 접수 안내", level=2)
p = doc.add_paragraph()
p.add_run("접수 기간: ").bold = True
p.add_run("2026년 9월 1일(화) ~ 2026년 9월 15일(화) 18:00까지")

p = doc.add_paragraph()
p.add_run("접수 방법: ").bold = True
p.add_run("학교 포털 공모전 신청 게시판을 통해 온라인 접수")

doc.add_heading("2. 결과 발표", level=2)
p = doc.add_paragraph()
p.add_run("1차 서류 심사 결과 발표: ").bold = True
p.add_run("2026년 9월 20일(일)")

p = doc.add_paragraph()
p.add_run("최종 결과 발표: ").bold = True
p.add_run("2026년 9월 25일(금)")

doc.add_heading("3. 시상식", level=2)
p = doc.add_paragraph()
p.add_run("일시: ").bold = True
p.add_run("2026년 10월 5일(월) 오후 2시")

p = doc.add_paragraph()
p.add_run("장소: ").bold = True
p.add_run("학생회관 대강당")

doc.add_heading("4. 문의", level=2)
doc.add_paragraph("창업지원단 (02-123-4567)")

out_path = r"C:\Users\12cog\timetable-recommender\scratchpad\sample_notice.docx"
doc.save(out_path)
print("saved:", out_path)
